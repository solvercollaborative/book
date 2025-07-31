from docx import Document
import os, re
import sys
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
import tempfile
import shutil

# --- E-Book Creation Functions Scaffold ---

def set_font_georgia(doc: Document) -> None:
    """Ensure all text in the document uses the 'Georgia' font, including all styles."""
    # Set font for all runs in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Georgia"
    # Set font for all runs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Georgia"
    # Set font for all runs in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Georgia"
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Georgia"
    # Set the font for all paragraph and character styles
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = "Georgia"

def adjust_title_page(doc: Document) -> None:
    """Change 'superArchItelligence' on the title page to size 28, 
        'Redesigning the real world' and 'for artificial intelligence' to size 20, 
        and ensure only one blank line above 'Alan G Street'. 
        Raise an error if 'superArchItelligence' is not found."""
    found_title = False
    # Change font size of 'superArchItelligence' on title page from 36 to 28
    for paragraph in doc.paragraphs[:10]:  # Assume title is in the first 10 paragraphs
        if "superArchItelligence" in paragraph.text:
            found_title = True
            for run in paragraph.runs:
                run.font.size = Pt(28)
        if "Redesigning the real world" in paragraph.text or "for artificial intelligence" in paragraph.text:
            for run in paragraph.runs:
                run.font.size = Pt(20)
    if not found_title:
        raise ValueError("'superArchItelligence' not found in the first 10 paragraphs of the document.")
    # Remove one blank line above 'Alan G Street'
    for i, paragraph in enumerate(doc.paragraphs):
        if "Alan G Street" in paragraph.text:
            if i > 1 and doc.paragraphs[i - 1].text == "" and doc.paragraphs[i - 2].text == "":
                p = doc.paragraphs[i - 2]._element
                p.getparent().remove(p)
            break

def adjust_copyright_page(doc: Document) -> None:
    """Remove blank line at the top of the copyright page."""
    for i, paragraph in enumerate(doc.paragraphs):
        if "Copyright © 2025" in paragraph.text:
            if i > 0 and doc.paragraphs[i - 1].text.strip() == "":
                p = doc.paragraphs[i - 1]._element
                p.getparent().remove(p)
            break

def convert_index_to_static_text(doc: Document) -> None:
    """Convert the index field to static text by clearing and replacing index paragraphs with their plain text content."""

    # Locate the Index section
    index_start = None
    for i, para in enumerate(doc.paragraphs):
        if is_index_heading(para):
            index_start = i
            break

    # Extract index paragraphs preserving their content
    index_paragraphs = []
    for i, para in enumerate(doc.paragraphs[index_start + 1:], start=index_start + 1):
        if para.text.strip().upper() in ["ACKNOWLEDGEMENTS", "ABOUT THE AUTHOR"]:
            break
        index_paragraphs.append((i, para.text))

    # Replace index field paragraphs with static versions
    for i, text in index_paragraphs:
        doc.paragraphs[i].clear()  # Clear all content (including field codes)
        doc.paragraphs[i].add_run(text)  # Add back the plain text content


# Helper function to detect index section
def is_index_heading(paragraph):
    return paragraph.text.strip().upper() == "INDEX"

def check_index_entries_single_page_number(doc: Document) -> None:
    """Ensure no index entry has more than one simple integer page number (after static conversion)."""
    # Locate the Index section
    index_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == "INDEX":
            index_start = i
            break
    if index_start is None:
        raise ValueError("No index section found in the document.")
    # Extract index paragraphs (static text)
    index_paragraphs = []
    for para in doc.paragraphs[index_start + 1:]:
        text = para.text.strip()
        if text.upper() in ["ACKNOWLEDGEMENTS", "ABOUT THE AUTHOR"]:
            break
        if text:
            index_paragraphs.append(text)
    # Check for multiple simple integer page numbers after static conversion
    for entry in index_paragraphs:
        if ',' in entry:
            # Get everything after the first comma
            after_first_comma = entry.split(',', 1)[1]
            # Remove parenthetical expressions (e.g., (1.4))
            after_first_comma = re.sub(r'\([^)]*\)', '', after_first_comma)
            # Find all simple integer tokens
            numbers = re.findall(r'\b\d+\b', after_first_comma)
            if len(numbers) > 1:
                raise ValueError(f"Index entry '{entry}' has more than one page number after static conversion.")

def sanitize_bookmark_name(name):
    import re
    # Bookmark names must start with a letter and contain only letters, numbers, and underscores
    name = re.sub(r'[^\w]', '_', name)
    if not name[0].isalpha():
        name = 'B_' + name
    return name[:40]  # Bookmark names must be ≤ 40 characters

def repackage_docx_from_dir(temp_dir, output_path):
    import zipfile
    import os
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
        for foldername, subfolders, filenames in os.walk(temp_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, temp_dir)
                print(f"Adding to zip: {arcname}")
                docx_zip.write(file_path, arcname)

def convert_xe_tags_to_bookmarks(docx_path, output_path):
    import os
    import shutil
    import re
    from zipfile import ZipFile
    from lxml import etree
    
    NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Dictionary to store mapping of index terms to bookmark names
    index_term_to_bookmark = {}
    # Dictionary to store mapping of bookmark names to their surrounding text
    bookmark_to_text = {}
    
    temp_dir = "temp_docx"
    try:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        
        with ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        
        parser = etree.XMLParser(ns_clean=True)
        with open(xml_path, 'rb') as f:
            tree = etree.parse(f, parser)
        
        root = tree.getroot()
        bookmark_id = 0
        
        # Find all XE-related field parts and collect them for deletion
        xe_runs_to_delete = set()
        
        # Find all instrText elements containing "XE" (complete or partial)
        instr_texts = root.xpath('//w:instrText[contains(text(), "XE")]', namespaces=NS)
        
        for instr in instr_texts:
            field_text = instr.text or ""
            
            # Try to extract index term from complete XE tags
            match = re.search(r'\s*XE "([^"]+)"', field_text)
            index_term = None
            if match:
                index_term = match.group(1)
            
            # Create a bookmark for this XE field (even if fragmented)
            run = instr.getparent()
            paragraph = run.getparent()
            
            # Try to find a nearby run with display text
            text_run = None
            surrounding_text = ""
            run_index = paragraph.index(run)
            
            # Look for display text in nearby runs (broader search)
            for i in range(max(0, run_index-5), min(len(paragraph), run_index+2)):
                check_run = paragraph[i]
                if check_run.tag.endswith('r'):
                    text_elements = check_run.xpath('.//w:t', namespaces=NS)
                    if text_elements:
                        run_text = ''.join([t.text or '' for t in text_elements])
                        surrounding_text += run_text + " "
                        if text_run is None and i < run_index:  # Use the last text run before XE
                            text_run = check_run
            
            surrounding_text = surrounding_text.strip()
            
            if text_run is not None:
                # Create bookmark with sequential name
                bookmark_name = f"xe_bookmark_{bookmark_id}"
                bookmark_start = etree.Element('{%s}bookmarkStart' % NS['w'])
                bookmark_start.set('{%s}id' % NS['w'], str(bookmark_id))
                bookmark_start.set('{%s}name' % NS['w'], bookmark_name)
                
                bookmark_end = etree.Element('{%s}bookmarkEnd' % NS['w'])
                bookmark_end.set('{%s}id' % NS['w'], str(bookmark_id))
                
                bookmark_id += 1
                
                # Store the mapping if we found an index term
                if index_term:
                    index_term_to_bookmark[index_term] = bookmark_name
                
                # Also store the surrounding text for fuzzy matching
                if surrounding_text:
                    bookmark_to_text[bookmark_name] = surrounding_text
                
                # Create point bookmark by placing start and end adjacent to each other
                text_run.addprevious(bookmark_start)
                text_run.addprevious(bookmark_end)  # Both before the text run, making them adjacent
            
            # Mark this run for deletion
            xe_runs_to_delete.add(run)
        
        # Also find any runs that contain XE field characters or related content
        # Look for fldChar elements that might be part of XE fields
        for fld_char in root.xpath('//w:fldChar', namespaces=NS):
            run = fld_char.getparent()
            # Check if this run or nearby runs contain XE-related content
            paragraph = run.getparent()
            run_index = paragraph.index(run)
            
            # Check runs around this fldChar for XE content
            for i in range(max(0, run_index-2), min(len(paragraph), run_index+3)):
                check_run = paragraph[i]
                if check_run.tag.endswith('r'):
                    # Check for XE in instrText or regular text
                    for text_elem in check_run.xpath('.//w:instrText | .//w:t', namespaces=NS):
                        if text_elem.text and 'XE' in text_elem.text:
                            xe_runs_to_delete.add(check_run)
                            break
        
        # Delete all identified XE-related runs
        for run in xe_runs_to_delete:
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)
        
        tree.write(xml_path, encoding='utf-8', xml_declaration=True)
        repackage_docx_from_dir(temp_dir, output_path)
        
        # Return both mappings
        return index_term_to_bookmark, bookmark_to_text
        
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

def verify_no_xe_tags(doc: Document) -> None:
    """Ensure no {XE} tags remain in the document."""
    pass

def link_index_entries_to_bookmarks(doc: Document, index_term_to_bookmark: dict, bookmark_to_text: dict) -> None:
    """Create hyperlinks from index entries to their corresponding bookmarks."""
    
    # Find the index section
    index_start = None
    for i, para in enumerate(doc.paragraphs):
        if is_index_heading(para):
            index_start = i
            break
    
    if index_start is None:
        print("Warning: No index section found")
        return
    
    # Process index paragraphs
    index_end = len(doc.paragraphs)
    for i in range(index_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.text.strip().upper() in ["ACKNOWLEDGEMENTS", "ABOUT THE AUTHOR"]:
            index_end = i
            break
    
    matched_count = 0
    
    # Go through each index paragraph and create hyperlinks
    for i in range(index_start + 1, index_end):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text:
            continue
            
        # Parse the index entry to extract the main term
        # Format is typically: "Term, page_number" or "Term, Subterm, page_number"
        if ',' in text:
            main_term = text.split(',')[0].strip()
        else:
            main_term = text.strip()
        
        # Skip single letters (section headers)
        if len(main_term) == 1:
            continue
        
        # Find corresponding bookmark
        bookmark_name = None
        
        # Method 1: Exact match with extracted index terms
        for index_term, bm_name in index_term_to_bookmark.items():
            if index_term == main_term or main_term in index_term:
                bookmark_name = bm_name
                break
        
        # Method 2: Fuzzy matching with surrounding text
        if not bookmark_name:
            best_match_score = 0
            best_bookmark = None
            
            for bm_name, surrounding_text in bookmark_to_text.items():
                # Calculate similarity score
                score = calculate_text_similarity(main_term, surrounding_text)
                if score > best_match_score and score > 0.25:  # Lowered threshold from 0.3 to 0.25
                    best_match_score = score
                    best_bookmark = bm_name
            
            if best_bookmark:
                bookmark_name = best_bookmark
        
        if bookmark_name:
            matched_count += 1
            # Clear the paragraph and recreate it with hyperlink
            original_text = para.text
            para.clear()
            
            # Split text into term part and page number part
            if ',' in original_text:
                parts = original_text.rsplit(',', 1)  # Split on last comma
                term_part = parts[0].strip()
                # page_part = parts[1].strip() if len(parts) > 1 else ""
                
                # Create hyperlink for the term part only (no page number)
                add_hyperlink_to_paragraph(para, term_part, bookmark_name)
                
                # Don't add the page number part back - remove this section:
                # if page_part:
                #     para.add_run(f", {page_part}")
            else:
                # No page number, just make the whole thing a hyperlink
                add_hyperlink_to_paragraph(para, original_text, bookmark_name)
        else:
            # For unlinked entries, also remove page numbers
            original_text = para.text
            if ',' in original_text:
                parts = original_text.rsplit(',', 1)  # Split on last comma
                term_part = parts[0].strip()
                # Replace the paragraph text with just the term (no page number)
                para.clear()
                para.add_run(term_part)
    
    print(f"Created hyperlinks for {matched_count} index entries")


def calculate_text_similarity(term, text):
    """Calculate a simple similarity score between an index term and surrounding text."""
    term_lower = term.lower()
    text_lower = text.lower()
    
    # Clean up the term for better matching
    # Remove common parenthetical content like (3.4.6), (1.1), etc.
    import re
    clean_term = re.sub(r'\s*\([^)]*\)\s*', ' ', term_lower).strip()
    
    # Remove common punctuation that might not appear in surrounding text
    clean_term = re.sub(r'[&/-]', ' ', clean_term)
    
    # Normalize whitespace
    clean_term = re.sub(r'\s+', ' ', clean_term)
    
    # Also clean the text similarly
    clean_text = re.sub(r'[&/-]', ' ', text_lower)
    clean_text = re.sub(r'\s+', ' ', clean_text)
    
    # Exact match on cleaned versions
    if clean_term in clean_text:
        return 1.0
    
    # Original exact match (keep for backwards compatibility)
    if term_lower in text_lower:
        return 1.0
    
    # Word overlap on cleaned versions
    term_words = set(clean_term.split())
    text_words = set(clean_text.split())
    
    # Remove very common words that don't help with matching
    common_words = {'the', 'and', 'or', 'of', 'in', 'to', 'for', 'with', 'a', 'an'}
    term_words = term_words - common_words
    text_words = text_words - common_words
    
    if not term_words:
        return 0.0
    
    overlap = len(term_words.intersection(text_words))
    base_score = overlap / len(term_words)
    
    # Boost score for abbreviations and acronyms
    # Check if any term words are acronyms that might appear expanded in text
    acronym_boost = 0
    for term_word in term_words:
        if len(term_word) <= 5 and term_word.isupper():
            # Look for potential expansion in text
            acronym_pattern = r'\b' + r'\w*\s+'.join(term_word.lower()) + r'\w*'
            if re.search(acronym_pattern, clean_text):
                acronym_boost += 0.3
    
    # Boost score for partial matches of multi-word terms
    partial_boost = 0
    if len(term_words) > 1:
        # Check if key words from the term appear in the text
        key_words = [w for w in term_words if len(w) > 3]  # Focus on longer words
        if key_words:
            key_overlap = len(set(key_words).intersection(text_words))
            if key_overlap > 0:
                partial_boost = (key_overlap / len(key_words)) * 0.2
    
    final_score = min(1.0, base_score + acronym_boost + partial_boost)
    return final_score

def add_hyperlink_to_paragraph(paragraph, text, bookmark_name):
    """Add a hyperlink to a bookmark within a paragraph."""
    from docx.oxml.shared import qn
    from docx.oxml import OxmlElement
    
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    # Create run for the hyperlink text
    run = OxmlElement('w:r')
    
    # Create run properties for hyperlink styling (blue, underlined)
    rPr = OxmlElement('w:rPr')
    
    # Set color to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    # Set underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    run.append(rPr)
    
    # Create text element
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)
    
    hyperlink.append(run)
    
    # Add hyperlink to paragraph
    paragraph._p.append(hyperlink)

def validate_index(doc: Document) -> None:
    """Ensure all index entries are linked and have no page numbers."""
    pass
# --- End E-Book Creation Functions Scaffold ---


def load_docx(filename: str) -> Document:
    """Load the DOCX file and validate it contains '8x10' and is a .docx file."""
    if not filename.lower().endswith('.docx'):
        raise ValueError("File must be a .docx file")
    
    if "8x10" not in filename:
        raise ValueError("Filename must contain '8x10'")
    
    doc = Document(filename)
    return doc


def main():
    """Main function to process the document."""
    import sys
    import os
    import tempfile
    from docx.shared import Pt
    
    if len(sys.argv) != 2:
        print("Usage: python create_ebook_from_print.py <input_filename>")
        sys.exit(1)
    
    filename = sys.argv[1]
    
    try:
        # Step 1: Load and validate the document
        doc = load_docx(filename)
        
        # Step 2: Set font to Georgia
        set_font_georgia(doc)
        
        # Step 3: Adjust title page
        adjust_title_page(doc)
        
        # Step 4: Adjust copyright page  
        adjust_copyright_page(doc)
        
        # Step 5: Convert index to static text
        convert_index_to_static_text(doc)
        
        # Step 6: Check index entries for single page numbers
        check_index_entries_single_page_number(doc)
        
        # Step 7: Create new filename for e-book version
        base_name = os.path.splitext(filename)[0]
        new_filename = base_name.replace("8x10", "e-book") + ".docx"
        
        # Save the document before XE tag conversion
        doc.save(new_filename)
        
        # Step 8: Convert XE tags to bookmarks (requires file manipulation)
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_filename = temp_file.name
        
        try:
            index_term_to_bookmark, bookmark_to_text = convert_xe_tags_to_bookmarks(new_filename, temp_filename)
            
            # Load the modified document and create hyperlinks
            doc_with_bookmarks = Document(temp_filename)
            link_index_entries_to_bookmarks(doc_with_bookmarks, index_term_to_bookmark, bookmark_to_text)
            
            # Save final version
            doc_with_bookmarks.save(new_filename)
            
            print(f"Cloned document saved as: {new_filename}")
            print(f"Created {len(index_term_to_bookmark)} exact index term mappings")
            print(f"Created {len(bookmark_to_text)} bookmark text mappings for fuzzy matching")
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_filename)
            except:
                pass
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()

