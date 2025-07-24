from docx import Document
import os, re
import sys
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
import tempfile
import shutil

# --- E-Book Creation Functions Scaffold ---

def set_font_georgia(doc: Document) -> None:
    """Ensure all text in the document uses the 'Georgia' font, including all styles."""
    # Set font for all runs in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Georgia"
    # Set font for all runs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Georgia"
    # Set font for all runs in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Georgia"
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Georgia"
    # Set the font for all paragraph and character styles
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = "Georgia"

def adjust_title_page(doc: Document) -> None:
    """Change 'superArchItelligence' on the title page to size 28, 
        'Redesigning the real world' and 'for artificial intelligence' to size 20, 
        and ensure only one blank line above 'Alan G Street'. 
        Raise an error if 'superArchItelligence' is not found."""
    found_title = False
    # Change font size of 'superArchItelligence' on title page from 36 to 28
    for paragraph in doc.paragraphs[:10]:  # Assume title is in the first 10 paragraphs
        if "superArchItelligence" in paragraph.text:
            found_title = True
            for run in paragraph.runs:
                run.font.size = Pt(28)
        if "Redesigning the real world" in paragraph.text or "for artificial intelligence" in paragraph.text:
            for run in paragraph.runs:
                run.font.size = Pt(20)
    if not found_title:
        raise ValueError("'superArchItelligence' not found in the first 10 paragraphs of the document.")
    # Remove one blank line above 'Alan G Street'
    for i, paragraph in enumerate(doc.paragraphs):
        if "Alan G Street" in paragraph.text:
            if i > 1 and doc.paragraphs[i - 1].text == "" and doc.paragraphs[i - 2].text == "":
                p = doc.paragraphs[i - 2]._element
                p.getparent().remove(p)
            break

def adjust_copyright_page(doc: Document) -> None:
    """Remove blank line at the top of the copyright page."""
    for i, paragraph in enumerate(doc.paragraphs):
        if "Copyright © 2025" in paragraph.text:
            if i > 0 and doc.paragraphs[i - 1].text.strip() == "":
                p = doc.paragraphs[i - 1]._element
                p.getparent().remove(p)
            break

def remove_page_numbers_from_toc_tof(doc: Document) -> None:
    """Toggle off page numbers in TOC and TOF by removing trailing numbers after tabs in those sections."""
    toc_headings = ["Table of Contents", "Contents"]
    tof_headings = ["Table of Figures", "List of Figures"]
    in_toc = False
    in_tof = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # Detect start of TOC or TOF
        if any(h in text for h in toc_headings):
            in_toc = True
            in_tof = False
            continue
        if any(h in text for h in tof_headings):
            in_tof = True
            in_toc = False
            continue
        # Detect end of TOC/TOF (next heading or empty line after section)
        if in_toc or in_tof:
            if text == "" or text.isupper():
                in_toc = False
                in_tof = False
                continue
            # Remove page number after tab (usually last tab-separated value)
            if "\t" in paragraph.text:
                parts = paragraph.text.rsplit("\t", 1)
                if len(parts) == 2 and parts[1].strip().isdigit():
                    # Remove the page number
                    paragraph.text = parts[0]

def convert_index_to_static_text(doc: Document) -> None:
    """Convert the index field to static text by clearing and replacing index paragraphs with their plain text content."""

    # Locate the Index section
    index_start = None
    for i, para in enumerate(doc.paragraphs):
        if is_index_heading(para):
            index_start = i
            break

    # Extract index paragraphs preserving their content
    index_paragraphs = []
    for i, para in enumerate(doc.paragraphs[index_start + 1:], start=index_start + 1):
        if para.text.strip().upper() in ["ACKNOWLEDGEMENTS", "ABOUT THE AUTHOR"]:
            break
        index_paragraphs.append((i, para.text))

    # Replace index field paragraphs with static versions
    for i, text in index_paragraphs:
        doc.paragraphs[i].clear()  # Clear all content (including field codes)
        doc.paragraphs[i].add_run(text)  # Add back the plain text content


# Helper function to detect index section
def is_index_heading(paragraph):
    return paragraph.text.strip().upper() == "INDEX"

def check_index_entries_single_page_number(doc: Document) -> None:
    """Ensure no index entry has more than one simple integer page number (after static conversion)."""
    # Locate the Index section
    index_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == "INDEX":
            index_start = i
            break
    if index_start is None:
        raise ValueError("No index section found in the document.")
    # Extract index paragraphs (static text)
    index_paragraphs = []
    for para in doc.paragraphs[index_start + 1:]:
        text = para.text.strip()
        if text.upper() in ["ACKNOWLEDGEMENTS", "ABOUT THE AUTHOR"]:
            break
        if text:
            index_paragraphs.append(text)
    # Check for multiple simple integer page numbers after static conversion
    for entry in index_paragraphs:
        if ',' in entry:
            # Get everything after the first comma
            after_first_comma = entry.split(',', 1)[1]
            # Remove parenthetical expressions (e.g., (1.4))
            after_first_comma = re.sub(r'\([^)]*\)', '', after_first_comma)
            # Find all simple integer tokens
            numbers = re.findall(r'\b\d+\b', after_first_comma)
            if len(numbers) > 1:
                raise ValueError(f"Index entry '{entry}' has more than one page number after static conversion.")

def sanitize_bookmark_name(name):
    import re
    # Bookmark names must start with a letter and contain only letters, numbers, and underscores
    name = re.sub(r'[^\w]', '_', name)
    if not name[0].isalpha():
        name = 'B_' + name
    return name[:40]  # Bookmark names must be ≤ 40 characters

def repackage_docx_from_dir(temp_dir, output_path):
    import zipfile
    import os
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
        for foldername, subfolders, filenames in os.walk(temp_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, temp_dir)
                print(f"Adding to zip: {arcname}")
                docx_zip.write(file_path, arcname)

def convert_xe_tags_to_bookmarks(docx_path, output_path):
    import os
    import shutil
    import re
    from zipfile import ZipFile
    from lxml import etree
    NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    temp_dir = "temp_docx"
    try:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        with ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        parser = etree.XMLParser(ns_clean=True)
        with open(xml_path, 'rb') as f:
            tree = etree.parse(f, parser)
        root = tree.getroot()
        bookmark_id = 0
        instr_texts = root.xpath('//w:instrText[contains(text(), "XE")]', namespaces=NS)
        for instr in instr_texts:
            field_text = instr.text
            match = re.search(r'XE\s+"([^"]+)"', field_text)
            if not match:
                continue
            index_term = match.group(1)
            bookmark_name = sanitize_bookmark_name(index_term)
            run = instr.getparent()
            paragraph = run.getparent()
            # Try to find a nearby run with display text
            text_run = None
            run_index = paragraph.index(run)
            for i in reversed(range(run_index)):
                prior = paragraph[i]
                if prior.tag.endswith('r') and prior.xpath('.//w:t', namespaces=NS):
                    text_run = prior
                    break
            if text_run is not None:
                bookmark_start = etree.Element('{%s}bookmarkStart' % NS['w'], id=str(bookmark_id), name=bookmark_name)
                bookmark_end = etree.Element('{%s}bookmarkEnd' % NS['w'], id=str(bookmark_id))
                bookmark_id += 1
                text_run.addprevious(bookmark_start)
                text_run.addnext(bookmark_end)
            # Remove field characters (instrText + surrounding runs)
            for part in ['instrText', 'fldChar']:
                for node in run.xpath(f'.//w:{part}', namespaces=NS):
                    parent = node.getparent()
                    if parent is not None:
                        parent.getparent().remove(parent)
        tree.write(xml_path, encoding='utf-8', xml_declaration=True)
        # Properly repackage the DOCX using ZipFile
        repackage_docx_from_dir(temp_dir, output_path)
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

def verify_no_xe_tags(doc: Document) -> None:
    """Ensure no {XE} tags remain in the document."""
    pass

def link_index_entries_to_bookmarks(doc: Document, index_entries: dict) -> None:
    """Create a new two-column index with hyperlinks to bookmarks for each entry."""
    # Find the index heading
    index_start = None
    for i, p in enumerate(doc.paragraphs):
        if "INDEX" in p.text.strip():
            index_start = i
            break
    if index_start is None:
        return
    # Create a section with two columns
    section = doc.sections[-1]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
    # Insert bookmarks alphabetically
    for entry in sorted(index_entries.keys(), key=lambda s: s.lower()):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(entry)
        run.font.name = 'Georgia'
        run.font.size = Pt(10)
        run.bold = False

def validate_index(doc: Document) -> None:
    """Ensure all index entries are linked and have no page numbers."""
    pass
# --- End E-Book Creation Functions Scaffold ---


def load_docx(filename: str) -> Document:
    if "8x10" not in filename:
        raise ValueError("Filename must contain '8x10'")
    doc = Document(filename)
    return doc


def main():
    if len(sys.argv) < 2:
        print("Usage: python book_loader.py <path_to_docx_file>")
        return
    filename = sys.argv[1]
    try:
        doc = load_docx(filename)
        new_filename = filename.replace("8x10", "e-book")
        set_font_georgia(doc)
        adjust_title_page(doc)
        adjust_copyright_page(doc)
        #remove_page_numbers_from_toc_tof(doc)
        convert_index_to_static_text(doc)
        check_index_entries_single_page_number(doc)
        # Save intermediate e-book before bookmark step
        import tempfile, os, shutil
        temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
        os.close(temp_fd)
        doc.save(temp_path)
        # Use a temporary file for the bookmark step
        convert_xe_tags_to_bookmarks(temp_path, new_filename)
        os.remove(temp_path)
        print(f"Cloned document saved as: {new_filename}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()

